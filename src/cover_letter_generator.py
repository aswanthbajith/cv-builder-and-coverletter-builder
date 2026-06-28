"""World-class cover letter generation in DOCX format.

Requirements addressed:
1. Human readability: excellent
2. AI-generated text detection resistance
3. Include relevant keywords naturally
4. Evidence-based claims, not generic assertions
5. Quantified achievements where possible
6. Technical depth highlighted
7. Measurable impact emphasized
8. Relevance to target job
9. Compelling narrative structure
10. Strong opening hook + specific closing
"""

import logging
import random
from pathlib import Path
from typing import Dict, Any, List
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from job_automation.config import load_config

logger = logging.getLogger(__name__)

# Opening hooks by job type — varied, non-formulaic
OPENING_HOOKS = {
    'hpc': [
        "Most students learn parallel computing from textbooks. I learned it by watching a 12-minute simulation crawl to 7 minutes after rewriting nested loops into vectorized NumPy — and then wondering how much further GPU offloading could take it.",
        "I spent last semester profiling memory bottlenecks in a 30,000-entity simulation. The experience taught me that performance optimization is less about heroics and more about understanding where the data actually flows.",
        "When I first ran my Python simulation on the university cluster, I was surprised that the slowest part wasn't the math — it was the I/O. That realization shifted how I think about HPC problems."
    ],
    'quantum': [
        "My first encounter with quantum computing wasn't in a lecture hall — it was debugging a VQE implementation at 2 AM, trying to understand why the ansatz wasn't converging. That frustration taught me more than any textbook could.",
        "I came to quantum computing from a mechanical engineering background, which means I approach qubit optimization the same way I approach thermal modeling: identify the constraints, then find the best path through the state space."
    ],
    'software': [
        "As a working student, I once spent a week tracking a bug that turned out to be a single missing await in an async function. The fix took 30 seconds; the lesson lasted much longer.",
        "I joined my current team as a working student with basic Python skills. Six months later, I was leading code reviews for the same modules I once struggled to understand."
    ],
    'research': [
        "I learned early in my B.Tech that the best research questions come from failed experiments. My thermal simulation project didn't just model battery heat — it identified a runaway risk that changed the design.",
        "The 18-page report I wrote on large-scale data analysis wasn't my first attempt. It was my third rewrite, after realizing that the data had a story I hadn't been listening to."
    ],
    'default': [
        "I came to computational science through an unusual path: a mechanical engineering degree, a working student role in e-commerce, and a master's program that forced me to bridge both worlds.",
        "My background is a mix of engineering fundamentals and software practice. I don't just write code — I understand the physics it's supposed to represent."
    ]
}

# Closing variations — specific, not generic
CLOSING_VARIATIONS = [
    "I would welcome the chance to discuss how my experience with {specific} could contribute to {company}'s work on {topic}. I have attached my CV and can provide references on request.",
    "If you're looking for someone who can {action} while understanding the {context} behind it, I'd like to talk. My CV is attached, and I am available for an interview at your convenience.",
    "I am genuinely interested in {company}'s approach to {topic}, and I believe my background in {specific} would be a good fit. I would welcome the opportunity to discuss this further."
]


class CoverLetterGenerator:
    """Generate world-class, human-sounding cover letters in DOCX."""
    
    def __init__(self):
        self.output_dir = Path(load_config().paths.generated_dir) / 'CoverLetters' / 'docx'
        self.output_dir.mkdir(parents=True, exist_ok=True)
    
    def generate(self, job: Dict[str, Any], profile: Dict[str, Any],
                 match_result: Any, filename_base: str) -> str:
        """Generate tailored DOCX cover letter."""
        
        doc = Document()
        
        # Clean formatting
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)
        
        # Build content
        job_text = ' '.join([
            str(job.get('job_title', '')),
            str(job.get('job_description', '')),
            str(job.get('required_skills', ''))
        ]).lower()
        
        self._add_header(doc, profile)
        self._add_date(doc)
        self._add_recipient(doc, job)
        self._add_salutation(doc, job)
        self._add_opening_hook(doc, job, profile, job_text)
        self._add_evidence_body(doc, job, profile, match_result, job_text)
        self._add_specific_closing(doc, job, profile, job_text)
        
        output_path = self.output_dir / f"{filename_base}_CoverLetter.docx"
        doc.save(output_path)
        
        logger.info(f"Generated cover letter: {output_path}")
        return str(output_path)
    
    def _add_header(self, doc: Document, profile: Dict) -> None:
        """Add candidate contact information — clean, minimal."""
        name = profile.get('name', 'Candidate Name')
        contact = profile.get('contact', {})
        
        p = doc.add_paragraph()
        run = p.add_run(name)
        run.bold = True
        run.font.size = Pt(14)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        parts = []
        for key in ['email', 'phone', 'location', 'linkedin']:
            val = contact.get(key, '')
            if val:
                parts.append(str(val))
        
        if parts:
            p = doc.add_paragraph(' | '.join(parts))
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    def _add_date(self, doc: Document) -> None:
        """Add current date in European format."""
        from datetime import datetime
        doc.add_paragraph(datetime.now().strftime('%d %B %Y'))
    
    def _add_recipient(self, doc: Document, job: Dict) -> None:
        """Add company and position information."""
        company = job.get('company', 'Hiring Team')
        title = job.get('job_title', 'Position')
        
        doc.add_paragraph()
        doc.add_paragraph(f"{company}")
        p = doc.add_paragraph(f"Re: Application for {title}")
        if p.runs:
            p.runs[0].bold = True
    
    def _add_salutation(self, doc: Document, job: Dict) -> None:
        """Add greeting — try to find specific team, fallback to generic."""
        company = job.get('company', 'your organization')
        title = job.get('job_title', 'the position')
        
        # Try to infer department from title
        dept = self._infer_department(title)
        if dept:
            doc.add_paragraph(f"Dear {company} {dept} Team,")
        else:
            doc.add_paragraph(f"Dear {company} Hiring Team,")
    
    def _infer_department(self, title: str) -> str:
        """Infer department from job title."""
        title_lower = title.lower()
        if any(k in title_lower for k in ['hpc', 'performance', 'parallel', 'cluster', 'gpu']):
            return 'HPC'
        elif any(k in title_lower for k in ['quantum', 'qiskit', 'cirq']):
            return 'Quantum Computing'
        elif any(k in title_lower for k in ['software', 'developer', 'engineer']):
            return 'Engineering'
        elif any(k in title_lower for k in ['research', 'scientist', 'phd']):
            return 'Research'
        elif any(k in title_lower for k in ['data', 'ml', 'machine learning']):
            return 'Data Science'
        elif any(k in title_lower for k in ['simulation', 'modeling']):
            return 'Simulation'
        return ''
    
    def _add_opening_hook(self, doc: Document, job: Dict, profile: Dict, job_text: str) -> None:
        """Add a compelling, non-formulaic opening paragraph."""
        
        # Determine job category for hook selection
        category = 'default'
        if any(k in job_text for k in ['hpc', 'high performance', 'parallel', 'mpi', 'gpu', 'cluster']):
            category = 'hpc'
        elif any(k in job_text for k in ['quantum', 'qiskit', 'cirq']):
            category = 'quantum'
        elif any(k in job_text for k in ['software', 'developer', 'web', 'frontend']):
            category = 'software'
        elif any(k in job_text for k in ['research', 'phd', 'scientist', 'thesis']):
            category = 'research'
        
        hooks = OPENING_HOOKS.get(category, OPENING_HOOKS['default'])
        hook = random.choice(hooks)
        
        doc.add_paragraph(hook)
        
        # Bridge to the specific role
        company = job.get('company', 'the organization')
        title = job.get('job_title', 'this position')
        
        bridge = (
            f"That is why I am applying for the {title} at {company}. "
            f"I believe the role aligns with both my current skills and the direction I want to grow in."
        )
        doc.add_paragraph(bridge)
    
    def _add_evidence_body(self, doc: Document, job: Dict, profile: Dict, 
                           match_result: Any, job_text: str) -> None:
        """Add evidence-based body paragraphs, not generic claims."""
        
        # Paragraph 1: Technical evidence — specific, quantified
        tech_evidence = self._build_technical_evidence(job, profile, job_text)
        doc.add_paragraph(tech_evidence)
        
        # Paragraph 2: Experience evidence — real outcomes
        exp_evidence = self._build_experience_evidence(job, profile, job_text)
        if exp_evidence:
            doc.add_paragraph(exp_evidence)
        
        # Paragraph 3: Project evidence — most relevant project
        proj_evidence = self._build_project_evidence(job, profile, job_text)
        if proj_evidence:
            doc.add_paragraph(proj_evidence)
        
        # Paragraph 4: Why this company specifically
        company_why = self._build_company_why(job, job_text, profile)
        doc.add_paragraph(company_why)
    
    def _build_technical_evidence(self, job: Dict, profile: Dict, job_text: str) -> str:
        """Build a technical evidence paragraph with specific tools and outcomes."""
        
        # Identify relevant skills from profile that match job
        skills = profile.get('technical_skills', {})
        if isinstance(skills, dict):
            all_skills = []
            for cat_skills in skills.values():
                all_skills.extend(cat_skills)
        else:
            all_skills = skills
        
        matched_skills = []
        for skill in all_skills:
            skill_lower = skill.lower()
            for part in skill_lower.split():
                if part in job_text and len(part) > 2:
                    matched_skills.append(skill)
                    break
        
        # Limit to 3-4 most relevant skills
        top_skills = matched_skills[:4] if matched_skills else ['Python', 'scientific computing']
        
        # Build evidence sentence
        if len(top_skills) >= 3:
            skills_str = ', '.join(top_skills[:-1]) + f', and {top_skills[-1]}'
        else:
            skills_str = ' and '.join(top_skills)
        
        # Add specific evidence
        evidence_parts = [f"My technical toolkit includes {skills_str}."]
        
        # Add quantified outcome if available
        achievements = profile.get('key_achievements', [])
        for ach in achievements:
            if any(k in ach.lower() for k in job_text.split()):
                evidence_parts.append(f"For example, {ach.lower()}.")
                break
        
        # Add learning/adaptability note
        evidence_parts.append(
            "I am comfortable picking up new tools quickly — my current role required me to learn the team's "
            "React/Vue stack in under two weeks to contribute to the next sprint."
        )
        
        return ' '.join(evidence_parts)
    
    def _build_experience_evidence(self, job: Dict, profile: Dict, job_text: str) -> str:
        """Build an experience evidence paragraph."""
        
        experience = profile.get('experience', [])
        if not experience:
            return ""
        
        exp = experience[0]
        title = exp.get('title', 'Working Student')
        company = exp.get('company', 'a company')
        bullets = exp.get('description', [])
        
        # Select 2 most relevant, quantified bullets
        quantified_bullets = [b for b in bullets if any(c.isdigit() for c in b)]
        if len(quantified_bullets) >= 2:
            evidence = f"At {company}, I {quantified_bullets[0].lower()} and {quantified_bullets[1].lower()}."
        elif quantified_bullets:
            evidence = f"At {company}, I {quantified_bullets[0].lower()}."
        else:
            evidence = f"At {company}, I {bullets[0].lower() if bullets else 'worked on software development projects'}."
        
        return evidence
    
    def _build_project_evidence(self, job: Dict, profile: Dict, job_text: str) -> str:
        """Build a project evidence paragraph."""
        
        projects = profile.get('projects', [])
        if not projects:
            return ""
        
        # Find most relevant project
        scored = []
        for project in projects:
            score = 0
            proj_text = str(project).lower()
            for word in job_text.split():
                if len(word) > 3 and word in proj_text:
                    score += 1
            scored.append((score, project))
        
        scored.sort(key=lambda x: x[0], reverse=True)
        if not scored or scored[0][0] == 0:
            return ""
        
        project = scored[0][1]
        name = project.get('name', 'a recent project')
        desc = project.get('description', [])
        
        if isinstance(desc, list) and desc:
            # Pick the most quantified bullet
            quantified = [b for b in desc if any(c.isdigit() for c in b)]
            if quantified:
                evidence = f"In {name}, I {quantified[0].lower()}."
            else:
                evidence = f"In {name}, I {desc[0].lower()}."
        elif desc:
            evidence = f"In {name}, I {str(desc).lower()}."
        else:
            return ""
        
        return evidence
    
    def _build_company_why(self, job: Dict, job_text: str, profile: Dict) -> str:
        """Build a paragraph explaining why this specific company."""
        
        company = job.get('company', 'the organization')
        
        # Extract specific themes from job description
        themes = []
        if 'quantum' in job_text:
            themes.append('quantum research')
        if 'hpc' in job_text or 'high performance' in job_text:
            themes.append('high-performance computing infrastructure')
        if 'machine learning' in job_text or 'ai' in job_text:
            themes.append('machine learning applications')
        if 'simulation' in job_text or 'modeling' in job_text:
            themes.append('computational modeling')
        if 'software' in job_text or 'developer' in job_text:
            themes.append('software engineering')
        if 'research' in job_text:
            themes.append('research environment')
        
        theme_str = themes[0] if themes else 'technical work'
        
        return (
            f"What draws me to {company} is the opportunity to work on {theme_str} "
            f"with a team that values both engineering depth and code quality. "
            f"I am ready to contribute from day one and learn what I do not yet know."
        )
    
    def _add_specific_closing(self, doc: Document, job: Dict, profile: Dict, job_text: str) -> None:
        """Add a closing that is specific, not generic."""
        
        company = job.get('company', 'the organization')
        title = job.get('job_title', 'the position')
        
        # Extract specific evidence to reference
        specific = "simulation optimization and software testing"
        if 'quantum' in job_text:
            specific = "quantum algorithms and numerical optimization"
        elif 'hpc' in job_text:
            specific = "parallel computing and performance profiling"
        elif 'software' in job_text:
            specific = "Python development and Agile workflows"
        elif 'machine learning' in job_text:
            specific = "machine learning pipelines and data analysis"
        
        topic = "computational challenges"
        if 'quantum' in job_text:
            topic = "quantum computing research"
        elif 'hpc' in job_text:
            topic = "high-performance computing systems"
        elif 'software' in job_text:
            topic = "software engineering problems"
        elif 'simulation' in job_text:
            topic = "simulation and modeling challenges"
        
        action = "write clean code and analyze complex datasets"
        if 'quantum' in job_text:
            action = "implement quantum algorithms and optimize numerical methods"
        elif 'hpc' in job_text:
            action = "profile performance bottlenecks and optimize parallel code"
        elif 'software' in job_text:
            action = "ship reliable features and maintain test coverage"
        
        context = "physics behind the computation"
        if 'software' in job_text:
            context = "engineering requirements driving the features"
        
        closing = random.choice(CLOSING_VARIATIONS).format(
            specific=specific,
            company=company,
            topic=topic,
            action=action,
            context=context
        )
        
        doc.add_paragraph(closing)
        
        # Sign-off
        doc.add_paragraph()
        doc.add_paragraph("Sincerely,")
        doc.add_paragraph()
        doc.add_paragraph(profile.get('name', 'Candidate Name'))
